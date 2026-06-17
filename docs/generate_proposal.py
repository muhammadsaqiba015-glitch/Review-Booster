#!/usr/bin/env python3
"""Generate an editable, polished Word (.docx) proposal for ReviewBoost (rateus.space)."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GREEN = RGBColor(0x05, 0x96, 0x69)
GREEN_DK = RGBColor(0x04, 0x78, 0x57)
DARK = RGBColor(0x0F, 0x17, 0x2A)
GREY = RGBColor(0x55, 0x5B, 0x66)
LIGHT_GREY = RGBColor(0x8A, 0x90, 0x9A)

doc = Document()
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor(0x26, 0x2A, 0x33)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.15


def shade(el, hex_fill):
    pr = el.get_or_add_tcPr() if el.tag.endswith('}tc') else None


def cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def no_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'none')
        borders.append(e)
    tblPr.append(borders)


def eyebrow(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text.upper()); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = GREEN


def heading(text, size=16):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(size); r.font.color.rgb = DARK
    return p


def body(text, color=None, size=11, after=8, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); r.font.size = Pt(size); r.bold = bold
    r.font.color.rgb = color or RGBColor(0x26, 0x2A, 0x33)
    return p


def step(num, lead, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    rn = p.add_run(f'{num}.  '); rn.bold = True; rn.font.color.rgb = GREEN; rn.font.size = Pt(11)
    rl = p.add_run(lead + ' '); rl.bold = True; rl.font.color.rgb = DARK
    p.add_run(text)


def bullet(lead, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(lead + ' '); r.bold = True; r.font.color.rgb = DARK
    p.add_run(text)


def callout(text):
    t = doc.add_table(rows=1, cols=1); no_borders(t)
    cell = t.cell(0, 0); cell_bg(cell, 'E8F8F1')
    cell.paragraphs[0].paragraph_format.space_before = Pt(4)
    cell.paragraphs[0].paragraph_format.space_after = Pt(4)
    r = cell.paragraphs[0].add_run(text); r.font.size = Pt(11); r.bold = True; r.font.color.rgb = GREEN_DK
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def screenshot(label):
    t = doc.add_table(rows=1, cols=1); no_borders(t)
    cell = t.cell(0, 0); cell_bg(cell, 'F1F2F4')
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(14)
    r = p.add_run(f'[  Insert screenshot:  {label}  ]')
    r.italic = True; r.font.size = Pt(10); r.font.color.rgb = LIGHT_GREY
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def link_line(label, url):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(label); r.bold = True; r.font.color.rgb = DARK
    ru = p.add_run(url); ru.font.color.rgb = GREEN_DK; ru.underline = True


# ───────────────────────── Cover ─────────────────────────
title = doc.add_paragraph(); title.paragraph_format.space_after = Pt(2)
r = title.add_run('ReviewBoost'); r.bold = True; r.font.size = Pt(30); r.font.color.rgb = DARK
sub = doc.add_paragraph(); sub.paragraph_format.space_after = Pt(16)
r = sub.add_run('More Google reviews. Higher ratings. Repeat customers.')
r.font.size = Pt(13); r.font.color.rgb = GREY

body('ReviewBoost helps your business grow with one simple, proven idea: reward customers for '
     'leaving a Google review. The result is more reviews, a stronger rating, and customers who '
     'keep coming back.', size=12, color=GREY, after=10)

callout('83% of customers read Google reviews before deciding where to go — and 94% choose the '
        'business with more and better ratings.')

body('If your competitor has 1,300 reviews and you have 400, most customers will pick them — no '
     'matter how good your food or service is. ReviewBoost closes that gap by turning every happy '
     'customer into a review.', after=10)

# ───────────────────────── How we work ─────────────────────────
eyebrow('How it works')
heading('A reward loop that runs itself')
body('Customers leave you a Google review and, in return, receive a discount on their next visit. '
     'The reviews drive your long-term growth on Google; the discount brings customers back in the '
     'short term. Best of all, it runs automatically — your staff barely lift a finger.')

# ───────────────────────── Step 1 ─────────────────────────
eyebrow('Step 1')
heading('Onboard your business in 2 minutes')
body('We have already created a login for your business. To set up a new business yourself, follow '
     'these steps:', color=GREY, after=8)
step('i', 'Go to', '')
link_line('', 'https://rateus.space/onboard')
step('ii', 'Add your business name.', '')
step('iii', 'Add your Google review link.', 'Not sure where to find it? This short video shows you how:')
link_line('', 'https://www.youtube.com/watch?v=zKpXIm0mTsc')
step('iv', 'Add your WhatsApp number', 'and tap Continue.')
step('v', 'Create a username and password.', 'Use these to log in from any device — at the counter '
     'or on your staff\'s phones — at https://rateus.space/login')
step('vi', 'You are taken to your Admin dashboard,', 'where you can redeem coupons, see all '
     'submissions, set your discount, and edit your review link any time.')
screenshot('Your Admin dashboard')

# ───────────────────────── Step 2 ─────────────────────────
eyebrow('Step 2')
heading('Put the discount in front of every customer')
step('i', 'Print and place your QR code.', 'Download it from your dashboard and place it on every '
     'table and at the counter, with a clear message: "Rate us on Google and get 15% off your next '
     'visit." A ready-to-print design is below.')
step('ii', 'Make sure customers know the offer.', 'A quick reminder from your staff — "Leave us a '
     'review and get a discount on your next visit" — dramatically increases how many customers take part.')
screenshot('Table-tent / QR poster design')

# ───────────────────────── Step 3 ─────────────────────────
eyebrow('Step 3')
heading('The customer journey — from review to reward')
step('i', 'Scan & enter details.', 'The customer scans the QR code and enters their name and '
     'WhatsApp number.')
step('ii', 'Leave a review.', 'They tap "Leave a Review", your Google review page opens in a new '
     'tab, and they post their review and take a screenshot.')
step('iii', 'Get the reward.', 'They upload the screenshot as proof. The coupon is issued instantly '
     '— shown on screen and sent to their WhatsApp so they always have it for their next visit. '
     'Every screenshot is saved to your dashboard so you can verify authenticity any time. '
     '(Automatic AI verification is coming soon.)')
step('iv', 'Redeem on the next visit.', 'When the customer returns, they show their coupon at billing. '
     'Your staff enter the code in the Admin dashboard to redeem it. The system instantly flags any '
     'coupon that is already used, expired, or invalid.')
step('v', 'Discount applied.', 'On successful redemption, the discount is applied to the total bill.')
screenshot('Redeem a coupon (Admin)')

# ───────────────────────── Pricing ─────────────────────────
eyebrow('Pricing')
heading('Simple, risk-free pricing')
pbox = doc.add_paragraph(); pbox.paragraph_format.space_after = Pt(4)
r = pbox.add_run('Rs. 50 '); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = GREEN
r2 = pbox.add_run('per redeemed coupon'); r2.font.size = Pt(13); r2.font.color.rgb = GREY
body('No setup fee. No monthly fee. You only pay when a customer actually returns and uses their '
     'discount — which means a review was earned and a repeat visit happened.', after=4)
body('(Pricing is a starting point and can be tailored to your business.)', size=9, color=LIGHT_GREY)

# ───────────────────────── Close ─────────────────────────
eyebrow('Get started')
heading('Ready to start a free pilot?')
body('We can set up your business this week and have reviews coming in within days. No commitment — '
     'if it does not bring you reviews, you pay nothing.', color=GREY)

doc.add_paragraph()
c = doc.add_paragraph()
rc = c.add_run('Contact:  '); rc.bold = True; rc.font.color.rgb = DARK
c.add_run('[Your name]   ·   [Your phone / WhatsApp]   ·   [your email]')

out = '/Users/saqib/Downloads/reviewboost/docs/ReviewBoost-Proposal.docx'
doc.save(out)
print('Saved:', out)
